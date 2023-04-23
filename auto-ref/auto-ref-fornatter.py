from typing import Generator
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Inches, Pt

"""
Don't ask me why I didn't create main()

This script only serve a purpose of automating the process of formatting the
References part or the Bibliography section of Resesarch paper.
This is a very basic script with lack of instructions given to us.

"""
# read files by line (just my habbit to use yield)
# It is assumed(or exactly what I did in my case) that
# each line in a file is the reference already in APA format


def read_lines(path: str) -> Generator:
    with open(path, "r") as f:
        for line in f.readlines():
            yield line.strip()

# Formating each references from few amd even vague instructions of adviser


def ref_format(ref: str, doc):
    # formats
    spacing = WD_LINE_SPACING.DOUBLE
    font_size = Pt(12)
    font_name = 'Arial'
    alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # objects
    paragraph = doc.add_paragraph()
    paragraph_format = paragraph.paragraph_format

    # paragraph formatting
    paragraph.alignment = alignment
    paragraph_format.line_spacing_rule = spacing
    paragraph_format.left_indent = Cm(1.27)  # add 0.5in indent to lines
    paragraph_format.first_line_indent = -Cm(1.27)  # Hanging indent format

    # write the reference continously
    run = paragraph.add_run(ref)

    # formatting run or the very text
    run.font.size = font_size
    run.font.name = font_name


# formats(basically in A4)
orientation = WD_ORIENTATION.PORTRAIT
height = Inches(11.7)
width = Inches(8.3)
margin = Inches(1)

# Read the references inside references.txt
references = list(read_lines("references.txt"))
references.sort()  # sorted it by author(given that a reference isn't sorted)
document = Document()

# apply the formats
section = document.sections[0]
section.orientation = orientation
section.page_height = height
section.page_width = width
section.left_margin = section.right_margin = section.top_margin = \
    section.bottom_margin = margin

# write the references in document
for reference in references:
    ref_format(reference, document)
document.save('references.docx')
